# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Định nghĩa màu sắc Trắng - Đen chuyên nghiệp
COLOR_PRIMARY = RGBColor(0, 0, 0)          # #000000 (Đen)
COLOR_SECONDARY = RGBColor(51, 51, 51)     # #333333 (Xám đậm)
COLOR_TEXT = RGBColor(51, 51, 51)          # #333333 (Xám đậm)
COLOR_MUTED = RGBColor(100, 100, 100)      # #646464 (Xám vừa)

HEX_BORDER = "7F8C8D"                      # Viền xám trung tính dễ chịu
HEX_LIGHT_GRAY = "F5F5F5"                  # Nền xen kẽ rất nhạt

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập padding cho ô (đơn vị dxa: 20 dxa = 1 pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Thiết lập viền cho ô"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for edge, val in borders.items():
        if val:
            node = OxmlElement(f'w:{edge}')
            node.set(qn('w:val'), val.get('val', 'single'))
            node.set(qn('w:sz'), str(val.get('sz', 4)))
            node.set(qn('w:space'), str(val.get('space', 0)))
            node.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{edge}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    """Đổ màu nền cho ô"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(run):
    """Thêm mã trường PAGE để tự động hiển thị số trang trong Word"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_total_page_number(run):
    """Thêm mã trường NUMPAGES để hiển thị tổng số trang"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def parse_markdown_text(paragraph, text, default_color=COLOR_TEXT, is_bullet=False):
    """Parse text chứa markdown cơ bản (**bold**, [text](url)) và add vào paragraph"""
    parts = re.split(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        elif part.startswith('[') and '](' in part:
            # Markdown Link: [text](url) -> Trích xuất text
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = COLOR_PRIMARY
                run.font.underline = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
        else:
            # Normal text
            run = paragraph.add_run(part)
            run.font.color.rgb = default_color
            run.font.name = 'Arial'
            run.font.size = Pt(10)

def main():
    md_path = "markdown/15-bien-ban-nghiem-thu-thanh-richland.md"
    docx_path = "markdown/15-bien-ban-nghiem-thu-thanh-richland.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    doc = Document()
    
    # Thiết lập margin trang chuẩn (1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Thiết lập header/footer
        section.different_first_page_header_footer = True
        
        # Header (từ trang 2 trở đi)
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = header_p.add_run("BIÊN BẢN NGHIỆM THU & BÀN GIAO CRM RICH LAND")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        # Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun1 = footer_p.add_run("Trang ")
        frun1.font.name = 'Arial'
        frun1.font.size = Pt(9)
        frun1.font.color.rgb = COLOR_MUTED
        add_page_number(footer_p.add_run())
        frun2 = footer_p.add_run(" / ")
        frun2.font.name = 'Arial'
        frun2.font.size = Pt(9)
        frun2.font.color.rgb = COLOR_MUTED
        add_total_page_number(footer_p.add_run())

    # Đọc nội dung Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Xử lý Table
        if line.startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # Kết thúc một bảng -> Render bảng đó vào DOCX
            in_table = False
            render_table(doc, table_rows)
            table_rows = []
            
        if not line:
            i += 1
            continue
            
        # Tiêu đề lớn nhất (H1)
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = 'Arial'
            run.font.color.rgb = COLOR_PRIMARY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Tiêu đề H2
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Arial'
            run.font.color.rgb = COLOR_PRIMARY
            
        # Tiêu đề H3
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.font.color.rgb = COLOR_SECONDARY
            
        # List items (Bullet points)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            content = line[2:]
            parse_markdown_text(p, content, COLOR_TEXT, is_bullet=True)
            
        # Metadata / Ghi chú in nghiêng
        elif line.startswith('*Tài liệu tham chiếu:*'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
            run.font.color.rgb = COLOR_MUTED
            
        # Đoạn văn thông thường
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_markdown_text(p, line, COLOR_TEXT)
            
        i += 1
        
    # Trường hợp file kết thúc bằng Table
    if in_table and table_rows:
        render_table(doc, table_rows)
        
    # Lưu file DOCX
    try:
        doc.save(docx_path)
        print(f"Successfully generated black-and-white word docx at {docx_path}")
    except PermissionError:
        fallback_path = docx_path.replace(".docx", "-v2.docx")
        doc.save(fallback_path)
        print(f"Permission denied on {docx_path} (likely open in Word). Saved copy to {fallback_path}")

def render_table(doc, raw_rows):
    """Phân tích các dòng markdown table và vẽ bảng trắng đen tối giản chuyên nghiệp"""
    rows_data = []
    for r in raw_rows:
        cols = [col.strip() for col in r.split('|')[1:-1]]
        if all(re.match(r'^[-:]+$', c) for c in cols if c):
            continue
        rows_data.append(cols)
        
    if not rows_data:
        return
        
    num_rows = len(rows_data)
    num_cols = len(rows_data[0])
    
    # Tạo table trong word
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Thiết lập độ rộng cột tối ưu cho Biên bản nghiệm thu (phát hiện bảng 6 cột hoặc 2 cột)
    col_widths = []
    if num_cols == 6:
        # Cột: STT (0.5"), Chức năng (1.7"), Kịch bản (1.8"), Kết quả mong đợi (1.8"), Trạng thái (0.9"), Người ký/Ghi chú (0.8")
        col_widths = [Inches(0.5), Inches(1.5), Inches(1.8), Inches(1.8), Inches(0.9), Inches(1.0)]
    elif num_cols == 3:
        col_widths = [Inches(1.0), Inches(2.0), Inches(4.5)]
    else:
        total_width = Inches(7.5)
        width_per_col = total_width / num_cols
        col_widths = [width_per_col] * num_cols

    # Định nghĩa kiểu viền đen mỏng
    border_thin = {'sz': 4, 'val': 'single', 'color': HEX_BORDER}
    border_thick = {'sz': 8, 'val': 'single', 'color': '000000'} # Viền đen đậm hơn
    
    for r_idx, row in enumerate(table.rows):
        if r_idx >= len(rows_data):
            break
            
        data_row = rows_data[r_idx]
        
        # Chiều cao hàng
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '400')
        trPr.append(trHeight)
        
        # Điền dữ liệu cho từng ô
        for c_idx, cell in enumerate(row.cells):
            if c_idx >= len(data_row):
                break
                
            cell.width = col_widths[c_idx]
            cell.text = "" # Xóa text mặc định
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            content = data_row[c_idx]
            
            # Cấu hình padding ô
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            if r_idx == 0:
                # Dòng đầu tiên (HEADER): KHÔNG ĐỔ MÀU NỀN (header ko cần background)
                # Chỉ vẽ viền đen đậm hơn ở dưới header để phân cách rõ ràng
                set_cell_borders(cell, top=border_thin, bottom=border_thick, left=border_thin, right=border_thin)
                
                run = p.add_run(content)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_PRIMARY # Chữ màu đen
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Dòng nội dung thường: Vẽ viền xám mỏng xung quanh
                set_cell_borders(cell, top=border_thin, bottom=border_thin, left=border_thin, right=border_thin)
                
                # Check nếu nội dung là STT hoặc Trạng thái trống thì căn giữa
                if num_cols == 6 and (c_idx == 0 or c_idx == 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                parse_markdown_text(p, content, COLOR_TEXT)

    # Thêm khoảng trống nhỏ sau bảng
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

if __name__ == '__main__':
    main()
